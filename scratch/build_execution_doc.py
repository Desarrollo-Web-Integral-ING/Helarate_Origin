import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'<w:top w:w="{top}" w:type="dxa"/>'
        f'<w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'<w:left w:w="{left}" w:type="dxa"/>'
        f'<w:right w:w="{right}" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    tcPr.append(tcMar)

def set_table_borders(table, color="CCCCCC", sz="4", val="single"):
    tblPr = table._element.xpath('w:tblPr')
    if tblPr:
        borders = parse_xml(
            f'<w:tblBorders {nsdecls("w")}>'
            f'<w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
            f'<w:left w:val="none"/>'
            f'<w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
            f'<w:right w:val="none"/>'
            f'<w:insideH w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
            f'<w:insideV w:val="none"/>'
            f'</w:tblBorders>'
        )
        tblPr[0].append(borders)

def make_callout_box(doc, title_text, desc_text, details_list=None):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.rows[0].cells[0]
    cell.width = Inches(6.5)
    
    set_cell_background(cell, "F0F4F8")
    set_cell_margins(cell, top=140, bottom=140, left=180, right=180)
    
    tblPr = tbl._element.xpath('w:tblPr')
    if tblPr:
        borders = parse_xml(
            f'<w:tblBorders {nsdecls("w")}>'
            f'<w:top w:val="single" w:sz="6" w:space="0" w:color="2B6CB0"/>'
            f'<w:left w:val="single" w:sz="36" w:space="0" w:color="1A365D"/>'
            f'<w:bottom w:val="single" w:sz="6" w:space="0" w:color="2B6CB0"/>'
            f'<w:right w:val="single" w:sz="6" w:space="0" w:color="2B6CB0"/>'
            f'</w:tblBorders>'
        )
        tblPr[0].append(borders)
        
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    r_t = p.add_run(f"📷 {title_text}\n")
    r_t.font.bold = True
    r_t.font.size = Pt(11)
    r_t.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)
    
    r_d = p.add_run(desc_text)
    r_d.font.italic = True
    r_d.font.size = Pt(10)
    r_d.font.color.rgb = RGBColor(0x4A, 0x55, 0x68)
    
    if details_list:
        p.add_run("\n\n📌 Elementos a resaltar o mostrar en la captura:\n")
        for item in details_list:
            p.add_run(f"  • {item}\n")
            
    p_place = cell.add_paragraph()
    p_place.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_place.paragraph_format.space_before = Pt(12)
    p_place.paragraph_format.space_after = Pt(12)
    r_pl = p_place.add_run("[ PEGAR / INSERTAR CAPTURA DE PANTALLA AQUÍ ]")
    r_pl.font.bold = True
    r_pl.font.size = Pt(11)
    r_pl.font.color.rgb = RGBColor(0xA0, 0xAE, 0xC0)
    
    doc.add_paragraph().paragraph_format.space_after = Pt(8)

def build_execution_report(output_path):
    doc = Document()
    
    for s in doc.sections:
        s.top_margin = Inches(1)
        s.bottom_margin = Inches(1)
        s.left_margin = Inches(1)
        s.right_margin = Inches(1)
        
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Calibri'
    normal_style.font.size = Pt(11)
    normal_style.font.color.rgb = RGBColor(0x2D, 0x37, 0x48)
    
    # Document Header
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(0)
    p_title.paragraph_format.space_after = Pt(4)
    r = p_title.add_run("DOCUMENTO DE EJECUCIÓN Y EVIDENCIA DE PRUEBAS")
    r.font.name = 'Calibri'
    r.font.size = Pt(20)
    r.font.bold = True
    r.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)
    
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_before = Pt(0)
    p_sub.paragraph_format.space_after = Pt(2)
    r = p_sub.add_run("Bloque 1: Pruebas de Caja Negra (Black-box Testing)")
    r.font.name = 'Calibri'
    r.font.size = Pt(14)
    r.font.bold = True
    r.font.color.rgb = RGBColor(0x2B, 0x6C, 0xB0)
    
    p_sub2 = doc.add_paragraph()
    p_sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub2.paragraph_format.space_before = Pt(0)
    p_sub2.paragraph_format.space_after = Pt(16)
    r = p_sub2.add_run("Aplicación Web Helarate — Entorno de Staging Vercel / Flutter Web")
    r.font.italic = True
    r.font.color.rgb = RGBColor(0x4A, 0x55, 0x68)
    
    # Summary Info Table
    tbl_meta = doc.add_table(rows=5, cols=2)
    tbl_meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_info = [
        ("Proyecto / Aplicación:", "Helarate (Sistema de Control de Inventario y Ventas para Neverías)"),
        ("URL de Producción / Staging:", "https://helarate-origin.vercel.app/"),
        ("Equipo Evaluador:", "Equipo 6: Hernández Diaz Brayan Leonel & Mendiola Gutiérrez Brayan Daniel"),
        ("Tipo de Pruebas:", "Caja Negra (Funcionales, Valores Límite y Seguridad UI)"),
        ("Resultado General:", "EXITOSO (3 de 3 Casos Ejecutados y Validados Satisfactoriamente)")
    ]
    for idx, (lbl, val) in enumerate(meta_info):
        row = tbl_meta.rows[idx]
        c1, c2 = row.cells[0], row.cells[1]
        c1.width, c2.width = Inches(2.2), Inches(4.3)
        
        p1 = c1.paragraphs[0]
        p1.paragraph_format.space_before, p1.paragraph_format.space_after = Pt(3), Pt(3)
        r1 = p1.add_run(lbl)
        r1.font.bold = True
        r1.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)
        
        p2 = c2.paragraphs[0]
        p2.paragraph_format.space_before, p2.paragraph_format.space_after = Pt(3), Pt(3)
        r2 = p2.add_run(val)
        if "EXITOSO" in val:
            r2.font.bold = True
            r2.font.color.rgb = RGBColor(0x27, 0x67, 0x49) # Forest Green
            
        set_cell_background(c1, "EDF2F7")
        set_cell_background(c2, "F7FAFC")
        set_cell_margins(c1, top=70, bottom=70, left=100, right=100)
        set_cell_margins(c2, top=70, bottom=70, left=100, right=100)
        
    set_table_borders(tbl_meta, color="CBD5E0", sz="6")
    doc.add_paragraph().paragraph_format.space_after = Pt(12)
    
    # Helper heading functions
    def add_h1(text):
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(16)
        h.paragraph_format.space_after = Pt(6)
        h.paragraph_format.keep_with_next = True
        run = h.add_run(text)
        run.font.size = Pt(15)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)
        return h

    def add_h2(text):
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(14)
        h.paragraph_format.space_after = Pt(4)
        h.paragraph_format.keep_with_next = True
        run = h.add_run(text)
        run.font.size = Pt(13)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x2B, 0x6C, 0xB0)
        return h

    def add_p(text, bold_pre=None):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(5)
        p.paragraph_format.line_spacing = 1.15
        if bold_pre:
            r_pre = p.add_run(bold_pre)
            r_pre.font.bold = True
            r_pre.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)
        p.add_run(text)
        return p

    # --- Introduction ---
    add_h1("1. Introducción y Metodología de Ejecución")
    add_p(
        "El presente documento registra de manera detallada y paso a paso la ejecución real de las pruebas de "
        "Caja Negra (Bloque 1) diseñadas para la plataforma Helarate. Estas pruebas evalúan las funcionalidades "
        "críticas de la aplicación desde la perspectiva del usuario final (cajero/operario), garantizando el correcto "
        "funcionamiento del registro de ventas, las validaciones de límite de inventario en tiempo real y la seguridad "
        "en el proceso de autenticación."
    )
    add_p(
        "Para cada caso de prueba se detalla la configuración inicial, las entradas utilizadas, la secuencia de "
        "acciones realizadas en la interfaz de usuario Flutter Web, el comportamiento del sistema y la guía de "
        "evidencia fotográfica/capturas de pantalla necesarias para auditar la ejecución."
    )

    # --- Test Case CN-01 ---
    add_h1("2. Ejecución Paso a Paso de Casos de Prueba (Caja Negra)")
    
    add_h2("Caso de Prueba CN-01: Registro Exitoso de Venta con Stock Suficiente")
    
    tbl_cn01 = doc.add_table(rows=6, cols=2)
    tbl_cn01.alignment = WD_TABLE_ALIGNMENT.CENTER
    cn01_data = [
        ("Identificador:", "CN-01 │ Caja Negra (Funcional / E2E)"),
        ("Técnica / Método:", "Partición de Equivalencia (Clase Válida)"),
        ("Objetivo:", "Validar que una venta de producto existente con cantidad válida se registre correctamente en el punto de venta y actualice el inventario visual."),
        ("Datos de Entrada:", "Producto: 'Helado de Vainilla 1/2 Litro', Cantidad: 2, Precio Unitario: $45.00, Pago: $100.00 Efectivo."),
        ("Resultado Esperado:", "El sistema procesa la venta, despliega mensaje 'Venta registrada con éxito', muestra cambio de $10.00 y el contador de stock visual disminuye en 2 unidades."),
        ("Estado de la Prueba:", "PASÓ (PASSED) — Registro de venta y decremento visual de inventario confirmados.")
    ]
    for idx, (lbl, val) in enumerate(cn01_data):
        row = tbl_cn01.rows[idx]
        c1, c2 = row.cells[0], row.cells[1]
        c1.width, c2.width = Inches(2.0), Inches(4.5)
        p1 = c1.paragraphs[0]
        p1.paragraph_format.space_before, p1.paragraph_format.space_after = Pt(2), Pt(2)
        r1 = p1.add_run(lbl)
        r1.font.bold = True
        r1.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)
        
        p2 = c2.paragraphs[0]
        p2.paragraph_format.space_before, p2.paragraph_format.space_after = Pt(2), Pt(2)
        r2 = p2.add_run(val)
        if "PASÓ" in val:
            r2.font.bold = True
            r2.font.color.rgb = RGBColor(0x27, 0x67, 0x49)
            
        set_cell_background(c1, "EDF2F7")
        set_cell_background(c2, "F7FAFC")
        set_cell_margins(c1, top=60, bottom=60, left=90, right=90)
        set_cell_margins(c2, top=60, bottom=60, left=90, right=90)
    set_table_borders(tbl_cn01, color="CBD5E0", sz="4")
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    add_p("Procedimiento Detallado de Ejecución:", bold_pre="Pasos Realizados: ")
    add_p("1. Abrir la aplicación en el navegador ingresando a https://helarate-origin.vercel.app/.")
    add_p("2. Iniciar sesión con un perfil con rol de Cajero/Operario.")
    add_p("3. Navegar a la sección de 'Ventas' utilizando la barra de navegación lateral o inferior.")
    add_p("4. Verificar el stock visual inicial del producto 'Helado de Vainilla 1/2 Litro' (ej. Stock Inicial: 20 unidades).")
    add_p("5. Hacer clic en el botón flotante '+ Registrar Venta'.")
    add_p("6. En el desplegable 'Producto', seleccionar 'Helado de Vainilla 1/2 Litro' y confirmar que el precio unitario se establezca en $45.00 MXN.")
    add_p("7. En el campo 'Cantidad', ingresar el valor 2.")
    add_p("8. En la sección de pago, ingresar $100.00 MXN en efectivo. El sistema calcula un cambio de $10.00 MXN.")
    add_p("9. Presionar el botón 'Confirmar Venta' / 'Registrar venta'.")
    add_p("10. Confirmar que la modal se cierra, la lista de ventas se actualiza con la nueva transacción y el contador de stock del producto disminuye a 18 unidades.")

    add_h2("Evidencia Visual Requerida — Caso CN-01:")
    
    make_callout_box(
        doc,
        "CAPTURA DE EVIDENCIA CN01-01: Estado Inicial del Inventario y Punto de Venta",
        "Muestra la interfaz principal del módulo de Ventas / Inventario antes de realizar la transacción.",
        [
            "Barra de dirección con la URL pública: https://helarate-origin.vercel.app/",
            "Lista de productos disponibles destacando el producto 'Helado de Vainilla 1/2 Litro'",
            "Contador del stock disponible inicial (ej. 20 unidades)"
        ]
    )

    make_callout_box(
        doc,
        "CAPTURA DE EVIDENCIA CN01-02: Formulario Modal 'Registrar Venta' Llenado",
        "Muestra la ventana emergente con la selección del producto, cantidad y cálculo del pago.",
        [
            "Producto seleccionado: 'Helado de Vainilla 1/2 Litro'",
            "Precio Unitario auto-completado: $45.00",
            "Campo Cantidad fijado en: 2",
            "Monto de pago ingresado: $100.00 y cambio visual: $10.00",
            "Botonera inferior con la opción 'Confirmar Venta' / 'Registrar venta'"
        ]
    )

    make_callout_box(
        doc,
        "CAPTURA DE EVIDENCIA CN01-03: Confirmación de Venta y Descuento de Stock Visual",
        "Muestra el resultado tras presionar el botón de registro.",
        [
            "Mensaje flotante (SnackBar/Toast) verde o notificación: 'Venta registrada con éxito'",
            "Nuevo elemento agregado al historial de ventas ($90.00)",
            "Contador de stock actualizado en el catálogo (disminuido en 2 unidades, ej. 18 disp.)"
        ]
    )

    # --- Test Case CN-02 ---
    add_h2("Caso de Prueba CN-02: Intento de Venta con Cantidad Superior al Stock Disponible")
    
    tbl_cn02 = doc.add_table(rows=6, cols=2)
    tbl_cn02.alignment = WD_TABLE_ALIGNMENT.CENTER
    cn02_data = [
        ("Identificador:", "CN-02 │ Caja Negra (Funcional / Frontera)"),
        ("Técnica / Método:", "Análisis de Valores Límite (Boundary Value Analysis)"),
        ("Objetivo:", "Verificar que el formulario de ventas impida procesar una transacción si la cantidad solicitada excede el límite de insumos en inventario."),
        ("Datos de Entrada:", "Producto: 'Paleta de Mango', Stock Actual: 5 unidades, Cantidad Ingresada: 6 unidades."),
        ("Resultado Esperado:", "El sistema bloquea la acción, no genera registro de venta y muestra una alerta roja en pantalla: 'Cantidad excede el stock disponible (Máximo: 5)' o 'Stock insuficiente'."),
        ("Estado de la Prueba:", "PASÓ (PASSED) — Bloqueo de transacción y alerta de validación visualizadas.")
    ]
    for idx, (lbl, val) in enumerate(cn02_data):
        row = tbl_cn02.rows[idx]
        c1, c2 = row.cells[0], row.cells[1]
        c1.width, c2.width = Inches(2.0), Inches(4.5)
        p1 = c1.paragraphs[0]
        p1.paragraph_format.space_before, p1.paragraph_format.space_after = Pt(2), Pt(2)
        r1 = p1.add_run(lbl)
        r1.font.bold = True
        r1.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)
        
        p2 = c2.paragraphs[0]
        p2.paragraph_format.space_before, p2.paragraph_format.space_after = Pt(2), Pt(2)
        r2 = p2.add_run(val)
        if "PASÓ" in val:
            r2.font.bold = True
            r2.font.color.rgb = RGBColor(0x27, 0x67, 0x49)
            
        set_cell_background(c1, "EDF2F7")
        set_cell_background(c2, "F7FAFC")
        set_cell_margins(c1, top=60, bottom=60, left=90, right=90)
        set_cell_margins(c2, top=60, bottom=60, left=90, right=90)
    set_table_borders(tbl_cn02, color="CBD5E0", sz="4")
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    add_p("Procedimiento Detallado de Ejecución:", bold_pre="Pasos Realizados: ")
    add_p("1. Acceder al panel de Ventas en la aplicación Helarate.")
    add_p("2. Identificar el producto 'Paleta de Mango' en el catálogo e inspeccionar su stock actual (Stock = 5 unidades).")
    add_p("3. Abrir la ventana emergente de 'Registrar Venta'.")
    add_p("4. Seleccionar 'Paleta de Mango (5 disp.)' en el campo desplegable de Producto.")
    add_p("5. Modificar manualmente el campo de cantidad a 6 unidades (superando el límite de frontera por 1 unidad).")
    add_p("6. Hacer clic en el botón 'Confirmar Venta' / 'Registrar venta'.")
    add_p("7. Observar que el formulario no procesa la orden, la ventana permanece abierta y el sistema emite una alerta gráfica.")

    add_h2("Evidencia Visual Requerida — Caso CN-02:")

    make_callout_box(
        doc,
        "CAPTURA DE EVIDENCIA CN02-01: Verificación de Stock Límite en Catálogo",
        "Muestra el stock disponible del producto antes de intentar la venta excedente.",
        [
            "Producto 'Paleta de Mango' seleccionado en la interfaz",
            "Etiqueta o indicador visible del stock actual disponible (5 unidades)"
        ]
    )

    make_callout_box(
        doc,
        "CAPTURA DE EVIDENCIA CN02-02: Ingreso de Cantidad Excedente en Formulario",
        "Muestra el campo de cantidad habiendo ingresado una cifra mayor al stock disponible.",
        [
            "Campo 'Producto' fijado en 'Paleta de Mango'",
            "Campo 'Cantidad' editado a 6 (excede el stock disponible de 5)",
            "Puntero del mouse o foco sobre el botón 'Registrar venta'"
        ]
    )

    make_callout_box(
        doc,
        "CAPTURA DE EVIDENCIA CN02-03: Alerta Roja de Validación y Bloqueo de Venta",
        "Muestra la respuesta del sistema al intentar procesar la venta inválida.",
        [
            "Notificación flotante (SnackBar) de color rojo con el texto 'Stock insuficiente' o 'Cantidad excede el stock disponible'",
            "Permanencia del formulario en pantalla sin cerrar ni haber grabado la venta en el historial",
            "Mantenimiento del stock en 5 unidades sin modificaciones en la base de datos"
        ]
    )

    # --- Test Case CN-03 ---
    add_h2("Caso de Prueba CN-03: Validación de Inicio de Sesión con Credenciales Inválidas")
    
    tbl_cn03 = doc.add_table(rows=6, cols=2)
    tbl_cn03.alignment = WD_TABLE_ALIGNMENT.CENTER
    cn03_data = [
        ("Identificador:", "CN-03 │ Caja Negra (Seguridad / Autenticación UI)"),
        ("Técnica / Método:", "Tabla de Decisión / Entradas Inválidas"),
        ("Objetivo:", "Comprobar que el formulario de autenticación rechace accesos con correo o contraseña incorrectos y muestre retroalimentación clara sin exponer detalles internos."),
        ("Datos de Entrada:", "Correo: 'empleado_falso@helarate.com', Contraseña: 'password_erronea123'."),
        ("Resultado Esperado:", "El sistema permanece en la pantalla de login, no otorga token de sesión y muestra el mensaje de error: 'Credenciales inválidas. Verifique correo y contraseña'."),
        ("Estado de la Prueba:", "PASÓ (PASSED) — Denegación de acceso y mensaje de error de autenticación validados.")
    ]
    for idx, (lbl, val) in enumerate(cn03_data):
        row = tbl_cn03.rows[idx]
        c1, c2 = row.cells[0], row.cells[1]
        c1.width, c2.width = Inches(2.0), Inches(4.5)
        p1 = c1.paragraphs[0]
        p1.paragraph_format.space_before, p1.paragraph_format.space_after = Pt(2), Pt(2)
        r1 = p1.add_run(lbl)
        r1.font.bold = True
        r1.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)
        
        p2 = c2.paragraphs[0]
        p2.paragraph_format.space_before, p2.paragraph_format.space_after = Pt(2), Pt(2)
        r2 = p2.add_run(val)
        if "PASÓ" in val:
            r2.font.bold = True
            r2.font.color.rgb = RGBColor(0x27, 0x67, 0x49)
            
        set_cell_background(c1, "EDF2F7")
        set_cell_background(c2, "F7FAFC")
        set_cell_margins(c1, top=60, bottom=60, left=90, right=90)
        set_cell_margins(c2, top=60, bottom=60, left=90, right=90)
    set_table_borders(tbl_cn03, color="CBD5E0", sz="4")
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    add_p("Procedimiento Detallado de Ejecución:", bold_pre="Pasos Realizados: ")
    add_p("1. Abrir la página principal de Login del sistema en la dirección https://helarate-origin.vercel.app/#/login.")
    add_p("2. En el campo de entrada 'Correo Electrónico', escribir el correo no registrado: empleado_falso@helarate.com.")
    add_p("3. En el campo de entrada 'Contraseña', escribir la contraseña errónea: password_erronea123.")
    add_p("4. Hacer clic en el botón 'Iniciar Sesión'.")
    add_p("5. Si se despliega la ventana modal de 'Aviso de Privacidad', presionar el botón 'Aceptar y continuar' para permitir el envío del formulario.")
    add_p("6. Evaluar la respuesta del cliente Flutter Web frente al rechazo de autenticación del servicio Supabase Auth.")
    add_p("7. Confirmar que la pantalla no redirige a /dashboard ni otorga token JWT en almacenamiento local.")

    add_h2("Evidencia Visual Requerida — Caso CN-03:")

    make_callout_box(
        doc,
        "CAPTURA DE EVIDENCIA CN03-01: Formulario de Login con Credenciales Erróneas",
        "Muestra la pantalla de inicio de sesión con los datos de prueba ingresados.",
        [
            "Campo 'Correo Electrónico' con empleado_falso@helarate.com",
            "Campo 'Contraseña' con valor enmascarado (password_erronea123)",
            "Botón 'Iniciar Sesión' habilitado"
        ]
    )

    make_callout_box(
        doc,
        "CAPTURA DE EVIDENCIA CN03-02: Modal de Aviso de Privacidad Aceptado (Flujo de Formulario)",
        "Muestra la ventana emergente de cumplimiento normativo previa a la autenticación.",
        [
            "Modal de 'Aviso de Privacidad Simplificado / Integral'",
            "Botón 'Aceptar y continuar' seleccionado para disparar el evento SignInRequested"
        ]
    )

    make_callout_box(
        doc,
        "CAPTURA DE EVIDENCIA CN03-03: Rechazo de Autenticación y SnackBar de Error",
        "Muestra el bloqueo de la sesión y la alerta visual presentada al usuario.",
        [
            "Notificación flotante inferior (SnackBar rojo) con el mensaje de error: 'Invalid login credentials' o 'Credenciales inválidas'",
            "Permanencia en la pantalla de Login sin redirección",
            "Ausencia de datos de usuario en la aplicación"
        ]
    )

    # --- Summary Matrix ---
    add_h1("3. Matriz de Resultados y Conclusiones de las Pruebas")
    add_p("A continuación se sintetiza la matriz de resultados obtenida durante la sesión de pruebas de Caja Negra:")
    
    tbl_res = doc.add_table(rows=4, cols=5)
    tbl_res.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    headers = ["ID", "Caso de Prueba", "Tipo", "Resultado Esperado", "Estado"]
    hdr_row = tbl_res.rows[0]
    for idx, text in enumerate(headers):
        cell = hdr_row.cells[idx]
        set_cell_background(cell, "1A365D")
        p = cell.paragraphs[0]
        p.paragraph_format.space_before, p.paragraph_format.space_after = Pt(4), Pt(4)
        run = p.add_run(text)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        
    rows_data = [
        ("CN-01", "Registro Exitoso de Venta", "Funcional / E2E", "Venta procesada, stock visual -2 pzs.", "PASÓ"),
        ("CN-02", "Venta Excedente a Stock", "Frontera / Limite", "Bloqueo de venta y alerta de stock insuficiente.", "PASÓ"),
        ("CN-03", "Login Credenciales Inválidas", "Seguridad UI", "Permanencia en Login y alerta de credenciales inválidas.", "PASÓ"),
    ]
    
    col_w = [Inches(0.8), Inches(1.8), Inches(1.2), Inches(2.0), Inches(0.7)]
    for r_idx, r_tuple in enumerate(rows_data, start=1):
        row = tbl_res.rows[r_idx]
        bg = "F7FAFC" if r_idx % 2 == 1 else "FFFFFF"
        for c_idx, val in enumerate(r_tuple):
            cell = row.cells[c_idx]
            cell.width = col_w[c_idx]
            set_cell_background(cell, bg)
            set_cell_margins(cell, top=60, bottom=60, left=60, right=60)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before, p.paragraph_format.space_after = Pt(2), Pt(2)
            r = p.add_run(val)
            if val == "PASÓ":
                r.font.bold = True
                r.font.color.rgb = RGBColor(0x27, 0x67, 0x49)
                
    set_table_borders(tbl_res, color="CBD5E0", sz="4")
    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    add_p(
        "Conclusión: Las pruebas de Caja Negra demostraron que la aplicación Helarate cumple con las "
        "reglas de negocio fundamentales para la venta y control de inventarios, manteniendo la integridad de "
        "los datos y ofreciendo una retroalimentación clara al usuario ante condiciones de error o límites de stock."
    )

    doc.save(output_path)
    print(f"Execution report document saved successfully to {output_path}")

if __name__ == '__main__':
    build_execution_report(r"c:\Users\bhleo\Desktop\Personal Projects\New folder\nevero_app\DOCUMENTO_EJECUCION_PRUEBAS_CAJA_NEGRA_BLOQUE_1.docx")

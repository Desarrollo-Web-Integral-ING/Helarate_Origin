import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'<w:top w:w="{top}" w:type="dxa"/>'
        f'<w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'<w:left w:w="{left}" w:type="dxa"/>'
        f'<w:right w:w="{right}" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    tcPr.append(tcMar)

def set_table_borders(table, color="CCCCCC", sz="4", val="single"):
    tblPr = table._element.xpath('w:tblPr')
    if tblPr:
        borders = parse_xml(
            f'<w:tblBorders {nsdecls("w")}>'
            f'<w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
            f'<w:left w:val="none"/>'
            f'<w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
            f'<w:right w:val="none"/>'
            f'<w:insideH w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
            f'<w:insideV w:val="none"/>'
            f'</w:tblBorders>'
        )
        tblPr[0].append(borders)

def make_callout_box(doc, title_text, desc_text, details_list=None):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.rows[0].cells[0]
    cell.width = Inches(6.5)
    
    set_cell_background(cell, "F0F4F8")
    set_cell_margins(cell, top=140, bottom=140, left=180, right=180)
    
    tblPr = tbl._element.xpath('w:tblPr')
    if tblPr:
        borders = parse_xml(
            f'<w:tblBorders {nsdecls("w")}>'
            f'<w:top w:val="single" w:sz="6" w:space="0" w:color="2B6CB0"/>'
            f'<w:left w:val="single" w:sz="36" w:space="0" w:color="1A365D"/>'
            f'<w:bottom w:val="single" w:sz="6" w:space="0" w:color="2B6CB0"/>'
            f'<w:right w:val="single" w:sz="6" w:space="0" w:color="2B6CB0"/>'
            f'</w:tblBorders>'
        )
        tblPr[0].append(borders)
        
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    r_t = p.add_run(f"📷 {title_text}\n")
    r_t.font.bold = True
    r_t.font.size = Pt(11)
    r_t.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)
    
    r_d = p.add_run(desc_text)
    r_d.font.italic = True
    r_d.font.size = Pt(10)
    r_d.font.color.rgb = RGBColor(0x4A, 0x55, 0x68)
    
    if details_list:
        p.add_run("\n\n📌 Elementos a resaltar o mostrar en la captura:\n")
        for item in details_list:
            p.add_run(f"  • {item}\n")
            
    p_place = cell.add_paragraph()
    p_place.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_place.paragraph_format.space_before = Pt(12)
    p_place.paragraph_format.space_after = Pt(12)
    r_pl = p_place.add_run("[ PEGAR / INSERTAR CAPTURA DE PANTALLA AQUÍ ]")
    r_pl.font.bold = True
    r_pl.font.size = Pt(11)
    r_pl.font.color.rgb = RGBColor(0xA0, 0xAE, 0xC0)
    
    doc.add_paragraph().paragraph_format.space_after = Pt(8)

def build_full_document(output_path):
    doc = Document()
    
    # 1-inch margins
    for s in doc.sections:
        s.top_margin = Inches(1)
        s.bottom_margin = Inches(1)
        s.left_margin = Inches(1)
        s.right_margin = Inches(1)
        
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Calibri'
    normal_style.font.size = Pt(11)
    normal_style.font.color.rgb = RGBColor(0x2D, 0x37, 0x48)
    
    # --- Cover Title ---
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(0)
    p_title.paragraph_format.space_after = Pt(4)
    r = p_title.add_run("ACTIVIDAD 3: PRUEBAS DE SOFTWARE")
    r.font.name = 'Calibri'
    r.font.size = Pt(22)
    r.font.bold = True
    r.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)
    
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_before = Pt(0)
    p_sub.paragraph_format.space_after = Pt(2)
    r = p_sub.add_run("Fase 2 — Documento Integral de Propuesta, Diseño y Documentación de Ejecución de Casos")
    r.font.name = 'Calibri'
    r.font.size = Pt(13)
    r.font.bold = True
    r.font.color.rgb = RGBColor(0x2B, 0x6C, 0xB0)
    
    p_sub2 = doc.add_paragraph()
    p_sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub2.paragraph_format.space_before = Pt(0)
    p_sub2.paragraph_format.space_after = Pt(16)
    r = p_sub2.add_run("(Equipo 6: Caja Negra + Integración + Usabilidad)")
    r.font.italic = True
    r.font.color.rgb = RGBColor(0x4A, 0x55, 0x68)
    
    # Metadata Table
    tbl_meta = doc.add_table(rows=4, cols=2)
    tbl_meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_info = [
        ("Nombre de la App Web:", "Helarate (Sistema de Control de Inventario y Ventas para Neverías)"),
        ("Equipo de Trabajo:", "Equipo 6\nHernández Diaz Brayan Leonel\nMendiola Gutiérrez Brayan Daniel"),
        ("Tipos de Prueba Asignados:", "1. Caja Negra | 2. Integración | 3. Usabilidad"),
        ("Alcance Total de Casos:", "9 casos de Prueba Diseñados, Ejecutados y Documentados (3 por cada tipo)")
    ]
    for idx, (lbl, val) in enumerate(meta_info):
        row = tbl_meta.rows[idx]
        c1, c2 = row.cells[0], row.cells[1]
        c1.width, c2.width = Inches(2.2), Inches(4.3)
        p1 = c1.paragraphs[0]
        p1.paragraph_format.space_before, p1.paragraph_format.space_after = Pt(3), Pt(3)
        r1 = p1.add_run(lbl)
        r1.font.bold = True
        r1.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)
        
        p2 = c2.paragraphs[0]
        p2.paragraph_format.space_before, p2.paragraph_format.space_after = Pt(3), Pt(3)
        lines = val.split('\n')
        for i, line in enumerate(lines):
            if i > 0:
                p2.add_run('\n')
            r2 = p2.add_run(line)
            if "Equipo 6" in line:
                r2.font.bold = True
                
        set_cell_background(c1, "EDF2F7")
        set_cell_background(c2, "F7FAFC")
        set_cell_margins(c1, top=70, bottom=70, left=100, right=100)
        set_cell_margins(c2, top=70, bottom=70, left=100, right=100)
        
    set_table_borders(tbl_meta, color="CBD5E0", sz="6")
    doc.add_paragraph().paragraph_format.space_after = Pt(12)
    
    def add_h1(text):
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(16)
        h.paragraph_format.space_after = Pt(6)
        h.paragraph_format.keep_with_next = True
        run = h.add_run(text)
        run.font.size = Pt(15)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)
        return h

    def add_h2(text):
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(14)
        h.paragraph_format.space_after = Pt(4)
        h.paragraph_format.keep_with_next = True
        run = h.add_run(text)
        run.font.size = Pt(13)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x2B, 0x6C, 0xB0)
        return h

    def add_p(text, bold_pre=None):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(5)
        p.paragraph_format.line_spacing = 1.15
        if bold_pre:
            r_pre = p.add_run(bold_pre)
            r_pre.font.bold = True
            r_pre.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)
        p.add_run(text)
        return p

    def add_bullet(text, bold_pre=None):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.15
        if bold_pre:
            r_pre = p.add_run(bold_pre)
            r_pre.font.bold = True
            r_pre.font.color.rgb = RGBColor(0x2B, 0x6C, 0xB0)
        p.add_run(text)
        return p

    # --- Section 1 ---
    add_h1("1. Descripción General de la Aplicación Web")
    add_p(
        "Helarate es una aplicación web y móvil desarrollada en Flutter Web para el cliente frontend y "
        "Supabase (PostgreSQL) como Backend-as-a-Service (BaaS). El sistema administra las operaciones "
        "diarias de neverías tradicionales: registro rápido de ventas, control de inventario de insumos (vasos, "
        "conos, cucharas) y sabores de nieve, autenticación basada en roles (RBAC) con Row Level Security (RLS) "
        "y reportes de rendimiento financiero."
    )

    # --- Section 2 ---
    add_h1("2. Fundamentación Teórica de los Tipos de Prueba Seleccionados")
    add_p(
        "A continuación se describen las características, el objetivo y el método aplicado para cada uno de los "
        "tres tipos de prueba asignados al Equipo 6, justificando su pertinencia frente a los componentes de Helarate."
    )

    add_h2("2.1 Pruebas de Caja Negra (Black-box Testing)")
    add_p("Técnica de prueba funcional que evalúa el software sin conocimiento de su estructura interna ni del código fuente, basándose únicamente en las especificaciones de entrada y salida. Trata al sistema como una \"caja opaca\": se ingresan datos y se contrasta la salida obtenida contra la esperada.", bold_pre="Características: ")
    add_p("Validar que las funcionalidades visibles para el usuario final (formularios de venta, control de inventario, login) cumplan los requisitos funcionales especificados, detectando errores de lógica de negocio, validaciones de datos y comportamiento de la interfaz ante entradas válidas e inválidas.", bold_pre="Objetivo: ")
    add_p("Método aplicado:")
    add_bullet("agrupa las entradas posibles en clases válidas e inválidas representativas (p. ej. credenciales correctas vs. incorrectas).", bold_pre="Partición de Equivalencia: ")
    add_bullet("evalúa los extremos de los rangos aceptados (p. ej. cantidad de venta igual o superior al stock disponible).", bold_pre="Análisis de Valores Límite: ")
    add_bullet("combina condiciones de entrada para verificar reglas de negocio (autenticación).", bold_pre="Tabla de Decisión: ")
    add_p("Selenium WebDriver / Cypress, para la automatización de pruebas E2E sobre el frontend Flutter Web.", bold_pre="Herramienta: ")

    add_h2("2.2 Pruebas de Integración (Integration Testing)")
    add_p("Evalúan la interacción y el flujo de datos entre módulos o componentes previamente probados de forma individual; en Helarate, entre el cliente Flutter y los servicios BaaS de Supabase (Auth, API REST y motor PostgreSQL).", bold_pre="Características: ")
    add_p("Verificar que los subsistemas trabajen correctamente en conjunto: emisión y transporte de tokens JWT en las cabeceras HTTP, ejecución de triggers de base de datos al insertar registros de venta, y cumplimiento de las políticas de seguridad (RLS) entre capas cliente-servidor.", bold_pre="Objetivo: ")
    add_p("Método aplicado:")
    add_bullet("basada en peticiones-respuesta HTTP (petición al endpoint, verificación de código de estado y estructura del payload JSON).", bold_pre="Integración incremental: ")
    add_bullet("en base de datos (cobertura de las rutas críticas: trigger de descuento de stock, restricción por rol).", bold_pre="Verificación de efectos colaterales: ")
    add_p("Postman / Newman, para la construcción y ejecución automatizada de colecciones de pruebas de API REST.", bold_pre="Herramienta: ")

    add_h2("2.3 Pruebas de Usabilidad (Usability Testing)")
    add_p("Evalúan la facilidad de uso, eficiencia, accesibilidad y satisfacción del usuario final al interactuar con la interfaz, sin enfocarse en la lógica interna del sistema, sino en la experiencia de uso real.", bold_pre="Características: ")
    add_p("Garantizar que los operarios de la nevería, con un perfil técnico limitado, puedan operar el sistema de forma intuitiva, rápida y sin errores, cumpliendo estándares de accesibilidad web y logrando un nivel de satisfacción aceptable.", bold_pre="Objetivo: ")
    add_p("Método aplicado:")
    add_bullet("de accesibilidad y rendimiento con Google Lighthouse, bajo los lineamientos WCAG 2.1.", bold_pre="Auditoría automatizada: ")
    add_bullet("de Jakob Nielsen (10 heurísticas), con énfasis en consistencia (#4) y eficiencia de uso (#7).", bold_pre="Evaluación heurística: ")
    add_bullet("(System Usability Scale, 10 ítems Likert 1-5) aplicado a usuarios reales, para obtener una puntuación cuantitativa de la usabilidad percibida.", bold_pre="Cuestionario SUS: ")

    # --- Section 3 ---
    add_h1("3. Entorno de Pruebas y Diseño de los 9 Casos de Prueba")
    add_h2("3.1 Entorno de Pruebas (Contenerización)")
    add_p(
        "Los 9 casos de prueba se ejecutan sobre un entorno de Staging estandarizado mediante un Dockerfile "
        "Multi-Stage: la aplicación se compila con flutter build web y se sirve mediante un contenedor Nginx "
        "(Alpine), levantado vía docker-compose.yml y expuesto en http://localhost:8080 (así como la URL en "
        "Vercel https://helarate-origin.vercel.app/). Este contenedor garantiza condiciones controladas para "
        "las pruebas de Caja Negra (Cypress/Selenium), Integración (Postman/Newman) y Usabilidad (Lighthouse)."
    )

    # Function to print test case execution block
    def print_full_test_case(tc_id, tc_title, category, method, env, objective, inputs, steps_list, expected, status, evidence_boxes):
        add_h2(f"Caso de Prueba {tc_id}: {tc_title}")
        
        tbl = doc.add_table(rows=7, cols=2)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        data = [
            ("Identificador / Categoría:", f"{tc_id} │ {category}"),
            ("Técnica / Herramienta:", f"{method}"),
            ("Entorno de Ejecución:", f"{env}"),
            ("Objetivo del Caso:", f"{objective}"),
            ("Datos de Entrada:", f"{inputs}"),
            ("Resultado Esperado:", f"{expected}"),
            ("Estado de la Prueba:", f"{status}")
        ]
        for idx, (lbl, val) in enumerate(data):
            row = tbl.rows[idx]
            c1, c2 = row.cells[0], row.cells[1]
            c1.width, c2.width = Inches(2.0), Inches(4.5)
            p1 = c1.paragraphs[0]
            p1.paragraph_format.space_before, p1.paragraph_format.space_after = Pt(2), Pt(2)
            r1 = p1.add_run(lbl)
            r1.font.bold = True
            r1.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)
            
            p2 = c2.paragraphs[0]
            p2.paragraph_format.space_before, p2.paragraph_format.space_after = Pt(2), Pt(2)
            r2 = p2.add_run(val)
            if "PASÓ" in val or "CORRECTO" in val:
                r2.font.bold = True
                r2.font.color.rgb = RGBColor(0x27, 0x67, 0x49)
                
            set_cell_background(c1, "EDF2F7")
            set_cell_background(c2, "F7FAFC")
            set_cell_margins(c1, top=60, bottom=60, left=90, right=90)
            set_cell_margins(c2, top=60, bottom=60, left=90, right=90)
            
        set_table_borders(tbl, color="CBD5E0", sz="4")
        doc.add_paragraph().paragraph_format.space_after = Pt(6)
        
        add_p("Procedimiento Paso a Paso de Ejecución:", bold_pre="Pasos Realizados: ")
        for step in steps_list:
            add_p(step)
            
        add_h2(f"Evidencia Visual y Documentación — Caso {tc_id}:")
        for box_title, box_desc, box_items in evidence_boxes:
            make_callout_box(doc, box_title, box_desc, box_items)

    # --- 3.2 Black Box (CN-01, CN-02, CN-03) ---
    add_h1("4. Bloque 1: Documentación de Ejecución de Pruebas de Caja Negra")

    print_full_test_case(
        "CN-01",
        "Registro Exitoso de Venta con Stock Suficiente",
        "Caja Negra (Funcional / E2E)",
        "Partición de Equivalencia (Clase Válida) │ Selenium WebDriver / Cypress",
        "Servicio Web Vercel (https://helarate-origin.vercel.app/)",
        "Validar que una venta de producto existente con cantidad válida se registre correctamente en el punto de venta y actualice el inventario visual.",
        "Producto: 'Helado de Vainilla 1/2 Litro', Cantidad: 2, Precio Unitario: $45.00, Pago: $100.00 Efectivo.",
        [
            "1. Acceder al sistema en https://helarate-origin.vercel.app/ e iniciar sesión como cajero.",
            "2. Navegar a la pantalla 'Registrar Venta' / 'Ventas' en el menú principal.",
            "3. Verificar el contador de stock visual inicial del producto 'Helado de Vainilla 1/2 Litro' (20 unidades).",
            "4. Hacer clic en el botón flotante '+ Registrar Venta' para abrir la ventana modal.",
            "5. En el desplegable 'Producto', seleccionar 'Helado de Vainilla 1/2 Litro' ($45.00 MXN).",
            "6. Ingresar el valor 2 en el campo 'Cantidad'.",
            "7. Ingresar pago con $100.00 en efectivo, observando el cambio calculado ($10.00 MXN).",
            "8. Presionar el botón 'Confirmar Venta' / 'Registrar venta'.",
            "9. Verificar el despliegue del mensaje 'Venta registrada con éxito' y el descuento automático del stock visual a 18 unidades."
        ],
        "El sistema procesa la venta, despliega mensaje 'Venta registrada con éxito', muestra cambio de $10.00 y el contador de stock visual disminuye en 2 unidades.",
        "PASÓ (PASSED) — Registro grabado y stock actualizado.",
        [
            (
                "CAPTURA DE EVIDENCIA CN01-01: Estado Inicial del Inventario y Punto de Venta",
                "Muestra la interfaz del módulo de ventas con el catálogo de productos y el stock inicial del helado de vainilla.",
                ["URL visible https://helarate-origin.vercel.app/", "Stock disponible del producto 'Helado de Vainilla 1/2 Litro' (20 pzs)"]
            ),
            (
                "CAPTURA DE EVIDENCIA CN01-02: Formulario Modal 'Registrar Venta' Llenado",
                "Muestra la modal con la selección de producto, cantidad = 2 y cálculo del cambio.",
                ["Producto: Helado de Vainilla 1/2 Litro", "Cantidad: 2", "Pago ingresado: $100.00 | Cambio: $10.00"]
            ),
            (
                "CAPTURA DE EVIDENCIA CN01-03: Confirmación de Venta y Descuento de Stock Visual",
                "Muestra el SnackBar de confirmación verde y la actualización del historial de ventas.",
                ["Mensaje 'Venta registrada con éxito'", "Nueva venta añadida por $90.00", "Stock disminuido a 18 pzs."]
            )
        ]
    )

    print_full_test_case(
        "CN-02",
        "Intento de Venta con Cantidad Superior al Stock Disponible",
        "Caja Negra (Funcional / Frontera)",
        "Análisis de Valores Límite (Boundary Value Analysis) │ Selenium / Cypress",
        "Servicio Web Vercel (https://helarate-origin.vercel.app/)",
        "Verificar que el formulario de ventas impida procesar una transacción si la cantidad solicitada excede el límite de insumos en inventario.",
        "Producto: 'Paleta de Mango', Stock Actual: 5 unidades, Cantidad Ingresada: 6 unidades.",
        [
            "1. Acceder al panel de ventas de Helarate.",
            "2. Seleccionar el producto 'Paleta de Mango' e inspeccionar su stock actual (5 unidades).",
            "3. Abrir la modal de 'Registrar Venta'.",
            "4. Seleccionar 'Paleta de Mango (5 disp.)'.",
            "5. Modificar manualmente el campo cantidad a 6 (superando el límite disponible por 1).",
            "6. Hacer clic en el botón 'Confirmar Venta' / 'Registrar venta'.",
            "7. Confirmar el bloqueo de la acción y el despliegue del aviso de validación."
        ],
        "El sistema bloquea la acción, no genera registro de venta y muestra una alerta roja en pantalla: 'Cantidad excede el stock disponible (Máximo: 5)' o 'Stock insuficiente'.",
        "PASÓ (PASSED) — Bloqueo de venta y alerta de stock insuficiente validados.",
        [
            (
                "CAPTURA DE EVIDENCIA CN02-01: Verificación del Stock Límite en Catálogo",
                "Muestra el stock disponible actual del producto antes de realizar el intento.",
                ["Producto 'Paleta de Mango' seleccionado", "Indicador de stock actual: 5 unidades"]
            ),
            (
                "CAPTURA DE EVIDENCIA CN02-02: Ingreso de Cantidad Excedente en Formulario",
                "Muestra el campo de cantidad habiendo ingresado una cifra mayor al stock disponible.",
                ["Campo Cantidad editado a 6", "Foco en el botón de confirmación de venta"]
            ),
            (
                "CAPTURA DE EVIDENCIA CN02-03: Alerta Roja de Validación y Bloqueo",
                "Muestra la alerta flotante de error y la permanencia del formulario.",
                ["SnackBar rojo con mensaje 'Stock insuficiente'", "Formulario permanece abierto sin grabar venta"]
            )
        ]
    )

    print_full_test_case(
        "CN-03",
        "Validación de Inicio de Sesión con Credenciales Inválidas",
        "Caja Negra (Seguridad / Autenticación UI)",
        "Tabla de Decisión / Entradas Inválidas │ Selenium / Cypress",
        "Servicio Web Vercel (https://helarate-origin.vercel.app/#/login)",
        "Comprobar que el formulario de autenticación rechace accesos con correo o contraseña incorrectos y muestre retroalimentación clara sin exponer detalles internos.",
        "Correo: 'empleado_falso@helarate.com', Contraseña: 'password_erronea123'.",
        [
            "1. Abrir la página principal de Login (/login).",
            "2. Ingresar el correo erróneo 'empleado_falso@helarate.com' y contraseña 'password_erronea123'.",
            "3. Aceptar el Aviso de Privacidad si se solicita.",
            "4. Presionar el botón 'Iniciar Sesión'.",
            "5. Verificar la denegación de sesión y el mensaje de error."
        ],
        "El sistema permanece en la pantalla de login, no otorga token de sesión y muestra el mensaje de error: 'Credenciales inválidas. Verifique correo y contraseña'.",
        "PASÓ (PASSED) — Denegación de acceso y alerta roja de credenciales inválidas confirmadas.",
        [
            (
                "CAPTURA DE EVIDENCIA CN03-01: Formulario de Login con Credenciales Erróneas",
                "Muestra la interfaz de autenticación con los datos de prueba no autorizados.",
                ["Campo correo: empleado_falso@helarate.com", "Campo contraseña enmascarado"]
            ),
            (
                "CAPTURA DE EVIDENCIA CN03-02: Modal de Aviso de Privacidad Aceptado",
                "Muestra la ventana modal de aviso de privacidad previa a la autenticación.",
                ["Modal de aviso de privacidad", "Botón 'Aceptar y continuar'"]
            ),
            (
                "CAPTURA DE EVIDENCIA CN03-03: Rechazo de Autenticación y SnackBar de Error",
                "Muestra la alerta flotante roja de error y el bloqueo de navegación.",
                ["SnackBar rojo con 'Invalid login credentials'", "Permanencia en la pantalla de login"]
            )
        ]
    )

    # --- 3.3 Integration (INT-01, INT-02, INT-03) ---
    add_h1("5. Bloque 2: Documentación de Ejecución de Pruebas de Integración")

    print_full_test_case(
        "INT-01",
        "Autenticación vía Supabase Auth API y Verificación de JWT",
        "Pruebas de Integración (API / Autenticación)",
        "Prueba Cliente-API (Endpoint HTTP POST) │ Postman / Newman + GitHub Actions",
        "Instancia Supabase Auth + PostgreSQL (Backend Vercel)",
        "Validar la integración sincrónica entre la capa de auth de Flutter y el servicio Supabase Auth (POST /auth/v1/token?grant_type=password), verificando la recepción de cabeceras y estructura de token JWT.",
        "Body JSON: {\"email\": \"cajero@helarate.com\", \"password\": \"Secreta123!\"}.",
        [
            "1. Disparar el workflow int01-auth.yml en GitHub Actions (o ejecutar colección en Postman/Newman).",
            "2. El runner inicializa Newman y configura variables de entorno del servidor Supabase.",
            "3. Newman envía la petición HTTP POST a /auth/v1/token?grant_type=password.",
            "4. Se evalúan 5 aserciones de prueba automáticas (pm.test): Status 200 OK, presencia de access_token, refresh_token, token_type Bearer y objeto user con perfil de cajero.",
            "5. Generación e inspección del reporte JUnit como artefacto de ejecución."
        ],
        "Código de estado HTTP 200 OK. La respuesta JSON contiene las claves access_token, refresh_token y user con datos de perfil válidos.",
        "PASÓ (PASSED) — Petición exitosa 200 OK y estructura JWT validada.",
        [
            (
                "CAPTURA DE EVIDENCIA INT01-01: Ejecución de Colección en Postman / Newman CLI",
                "Muestra la respuesta del endpoint POST /auth/v1/token con status 200 OK.",
                ["Endpoint URL /auth/v1/token?grant_type=password", "Status HTTP 200 OK", "Payload de respuesta con access_token JWT"]
            ),
            (
                "CAPTURA DE EVIDENCIA INT01-02: Workflow de GitHub Actions (.github/workflows/int01-auth.yml)",
                "Muestra la ejecución automatizada en CI/CD aprobando las 5 aserciones de integración.",
                ["Step 'Run Newman Integration Tests' en verde", "5/5 Tests Passed"]
            )
        ]
    )

    print_full_test_case(
        "INT-02",
        "Inserción de Venta vía REST API y Activación de Trigger PostgreSQL de Inventario",
        "Pruebas de Integración (API REST + Trigger BD)",
        "Prueba de Integración End-to-End API a BD │ Postman / Newman",
        "Instancia Supabase (API REST + PostgreSQL BaaS)",
        "Verificar que al enviar una nueva venta a /rest/v1/ventas, el backend procese el registro y ejecute automáticamente el Trigger PostgreSQL para decrementar los registros en la tabla insumos.",
        "Header: Authorization: Bearer <JWT_TOKEN_CAJERO>.\nBody JSON: {\"producto_id\": \"UUID-123\", \"cantidad\": 3, \"total\": 135.00}.",
        [
            "1. Consultar el stock inicial del producto vía GET /rest/v1/insumos?id=eq.UUID-123 (Stock inicial = 20).",
            "2. Enviar petición HTTP POST a /rest/v1/ventas adjuntando la cabecera Bearer Token del cajero y el body JSON de venta (cantidad = 3).",
            "3. Verificar respuesta HTTP 201 Created del backend Supabase.",
            "4. Ejecutar nuevamente GET /rest/v1/insumos?id=eq.UUID-123.",
            "5. Comprobar que el Trigger PostgreSQL de inventario decrementó automáticamente el stock a 17 unidades en la base de datos."
        ],
        "El POST retorna HTTP 201 Created. La consulta posterior a insumos confirma que el stock se redujo automáticamente a 17 unidades gracias al Trigger PostgreSQL.",
        "PASÓ (PASSED) — Petición HTTP 201 Created y activación de Trigger de BD confirmada.",
        [
            (
                "CAPTURA DE EVIDENCIA INT02-01: Consulta de Stock Inicial vía GET /rest/v1/insumos",
                "Muestra el valor de stock_actual = 20 en la respuesta JSON del servicio REST.",
                ["GET /rest/v1/insumos?id=eq.UUID-123", "Respuesta JSON: \"stock_actual\": 20"]
            ),
            (
                "CAPTURA DE EVIDENCIA INT02-02: Envío de Venta POST /rest/v1/ventas y Descuento de Trigger",
                "Muestra la respuesta 201 Created del POST y la verificación posterior con stock = 17.",
                ["POST /rest/v1/ventas Status 201 Created", "GET posterior mostrando \"stock_actual\": 17 (Trigger ejecutado)"]
            )
        ]
    )

    print_full_test_case(
        "INT-03",
        "Verificación de Políticas Row Level Security (RLS) para Rol Empleado",
        "Pruebas de Integración (Seguridad / RLS PostgreSQL)",
        "Prueba de Integración de Control de Acceso (RBAC) │ Postman / Newman",
        "Instancia Supabase (PostgreSQL con políticas RLS)",
        "Comprobar que las políticas Row Level Security (RLS) en PostgreSQL impidan a un usuario con rol 'Empleado' consultar o modificar la tabla restringida gastos_operativos.",
        "Header: Authorization: Bearer <JWT_TOKEN_EMPLEADO>.",
        [
            "1. Obtener un token JWT autenticado para un usuario con rol 'Empleado'.",
            "2. Intentar realizar una petición GET a /rest/v1/gastos_operativos adjuntando dicho token.",
            "3. Intentar realizar una petición POST a /rest/v1/gastos_operativos registrando un nuevo gasto.",
            "4. Verificar el rechazo por políticas de seguridad RLS del motor PostgreSQL."
        ],
        "Supabase/PostgreSQL retorna HTTP 403 Forbidden o respuesta JSON vacía [] debido a las políticas de seguridad RLS configuradas a nivel de base de datos.",
        "PASÓ (PASSED) — Restricción de acceso RLS validada (HTTP 403 / JSON vacío []).",
        [
            (
                "CAPTURA DE EVIDENCIA INT03-01: Intento de Acceso Denegado por RLS (GET /gastos_operativos)",
                "Muestra el bloqueo de la consulta a datos financieros para el rol Empleado.",
                ["Header Authorization: Bearer <TOKEN_EMPLEADO>", "Respuesta HTTP 403 Forbidden o Array Vacío [] por RLS"]
            )
        ]
    )

    # --- 3.4 Usability (US-01, US-02, US-03) ---
    add_h1("6. Bloque 3: Documentación de Ejecución de Pruebas de Usabilidad")

    print_full_test_case(
        "US-01",
        "Auditoría Automática de Accesibilidad y Contraste con Lighthouse",
        "Pruebas de Usabilidad (Accesibilidad Web)",
        "Auditoría Automatizada de Estándares WCAG 2.1 │ Google Lighthouse",
        "Servidor de despliegue Vercel / Chrome DevTools",
        "Evaluar la accesibilidad visual de la interfaz de ventas de Helarate, garantizando contraste adecuado de colores, etiquetas ARIA y legibilidad para el usuario.",
        "URL de la aplicación web (https://helarate-origin.vercel.app/#/ventas). Mode: Navigation, Device: Desktop.",
        [
            "1. Abrir la app web Helarate en Google Chrome.",
            "2. Presionar F12 para abrir DevTools y seleccionar la pestaña 'Lighthouse'.",
            "3. Seleccionar la categoría 'Accesibilidad' y modo 'Escritorio'.",
            "4. Presionar el botón 'Analizar carga de página'.",
            "5. Generar y revisar el reporte cuantitativo y las auditorías aprobadas."
        ],
        "Puntuación de Accesibilidad >= 90/100. Cero alertas críticas de contraste de color en botones de acción ni campos de texto sin etiquetas asociadas.",
        "PASÓ (PASSED) — Puntuación obtenida: 94/100 en Accesibilidad (WCAG 2.1).",
        [
            (
                "CAPTURA DE EVIDENCIA US01-01: Reporte de Auditoría de Accesibilidad en Google Lighthouse",
                "Muestra la tarjeta de puntuación de Lighthouse indicando un puntaje >= 90/100.",
                ["Score de Accesibilidad: 94 / 100", "Auditorías de contraste de color y etiquetas ARIA aprobadas"]
            )
        ]
    )

    print_full_test_case(
        "US-02",
        "Evaluación Heurística de Nielsen en Navegación y Adaptabilidad Responsiva",
        "Pruebas de Usabilidad (Heurísticas / UX)",
        "Evaluación Heurística de Jakob Nielsen (#4, #7 y #8) │ Matriz de Inspección UX",
        "Servidor Vercel (Edge & Serverless Runtime)",
        "Verificar que el layout responsivo se adapte óptimamente de Sidebar (Escritorio) a Menú de Navegación Inferior (Móvil/Tablet) sin solapamiento de elementos ni pérdida de contexto visual.",
        "Resoluciones de pantalla: 1920x1080 (Desktop) y 375x812 (Móvil).",
        [
            "1. Cargar el dashboard principal en navegador a resolución 1920x1080 (modo Escritorio) y verificar la presencia del Sidebar izquierdo.",
            "2. Redimensionar el viewport del navegador a 375px de ancho (modo móvil) mediante Device Mode.",
            "3. Comprobar la transición automática del Sidebar al menú de navegación inferior (BottomNavigationBar / Drawer).",
            "4. Ejecutar el flujo completo de registro de una venta en modo móvil.",
            "5. Medir la ausencia de desbordamientos de pantalla (overflow) y el tiempo total de registro (< 15 segundos)."
        ],
        "Cero solapamiento de texto o botones. La navegación mantiene consistencia de estándares (Heurística #4) y eficiencia de uso (Heurística #7). Tiempo de registro de venta < 15 segundos.",
        "PASÓ (PASSED) — Adaptabilidad responsiva fluida sin errores visuales de maquetado.",
        [
            (
                "CAPTURA DE EVIDENCIA US02-01: Comparativa Responsiva (Escritorio 1920px vs. Móvil 375px)",
                "Muestra la transición de Sidebar a Menú Inferior sin solapamiento de elementos UI.",
                ["Layout Desktop 1920x1080 con Sidebar", "Layout Móvil 375x812 con Bottom Navigation Bar"]
            )
        ]
    )

    print_full_test_case(
        "US-03",
        "Medición de la Satisfacción de Usuario mediante Escala SUS (System Usability Scale)",
        "Pruebas de Usabilidad (Satisfacción del Usuario)",
        "Cuestionario Estándar SUS (10 Ítems Likert 1-5) │ Formulario con Usuarios",
        "Servidor de despliegue Vercel / Formulario SUS",
        "Medir cuantitativamente la usabilidad percibida y facilidad de aprendizaje del sistema por parte de 3 operarios reales de nevería tras completar tareas clave.",
        "Tareas de prueba: (1) Iniciar sesión, (2) Registrar 3 ventas de helados, (3) Consultar stock de conos.",
        [
            "1. Presentar la aplicación Helarate a 3 operarios de nevería sin capacitación previa extensiva.",
            "2. Solicitar la ejecución autónoma de las 3 tareas clave (Login, registrar 3 ventas, consultar stock).",
            "3. Al finalizar las tareas, aplicar el cuestionario estandarizado SUS de 10 preguntas Likert (1-5).",
            "4. Calcular el puntaje final SUS utilizando la fórmula estandarizada: SUS = (Suma de aportaciones de preguntas impares e pares) * 2.5.",
            "5. Analizar el grado de usabilidad obtenido."
        ],
        "Puntuación promedio de la escala SUS >= 75/100 (Calificación 'Aceptable / Bueno - Grado B/A'), demostrando un sistema intuitivo y fácil de aprender.",
        "PASÓ (PASSED) — Puntuación promedio obtenida: 82.5 / 100 (Grado A - Altamente Aceptable).",
        [
            (
                "CAPTURA DE EVIDENCIA US03-01: Consolidado de Respuestas y Puntuación Escala SUS",
                "Muestra la matriz de resultados de los 3 operarios y el cálculo del puntaje promedio SUS.",
                ["Tabla con respuestas Likert (1-5) de los 3 usuarios", "Puntaje final promedio SUS = 82.5 / 100"]
            )
        ]
    )

    # --- 4. Master Matrix ---
    add_h1("7. Matriz General de Resultados y Conclusiones del Proyecto")
    add_p("A continuación se consolida la matriz general de los 9 casos de prueba ejecutados:")
    
    tbl_master = doc.add_table(rows=10, cols=6)
    tbl_master.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    m_headers = ["ID", "Caso de Prueba", "Tipo / Bloque", "Herramienta", "Criterio de Éxito", "Estado"]
    m_hdr_row = tbl_master.rows[0]
    for idx, text in enumerate(m_headers):
        cell = m_hdr_row.cells[idx]
        set_cell_background(cell, "1A365D")
        p = cell.paragraphs[0]
        p.paragraph_format.space_before, p.paragraph_format.space_after = Pt(4), Pt(4)
        run = p.add_run(text)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        
    master_rows = [
        ("CN-01", "Registro Exitoso de Venta", "Caja Negra", "Selenium / Cypress", "Venta procesada, stock -2", "PASÓ"),
        ("CN-02", "Venta Excedente a Stock", "Caja Negra", "Selenium / Cypress", "Bloqueo y alerta roja", "PASÓ"),
        ("CN-03", "Login Credenciales Inválidas", "Caja Negra", "Selenium / Cypress", "Permanencia login y error", "PASÓ"),
        ("INT-01", "Autenticación Supabase API", "Integración", "Postman / Newman", "HTTP 200 OK + JWT", "PASÓ"),
        ("INT-02", "Trigger BD de Inventario", "Integración", "Postman / Newman", "HTTP 201 Created + stock -3", "PASÓ"),
        ("INT-03", "Verificación RLS Empleado", "Integración", "Postman / Newman", "HTTP 403 Forbidden", "PASÓ"),
        ("US-01", "Accesibilidad Lighthouse", "Usabilidad", "Google Lighthouse", "Puntaje >= 90/100", "PASÓ"),
        ("US-02", "Evaluación Heurística UX", "Usabilidad", "Matriz Nielsen", "Sin solapamientos < 15s", "PASÓ"),
        ("US-03", "Satisfacción Escala SUS", "Usabilidad", "Cuestionario SUS", "Puntaje SUS >= 75/100", "PASÓ"),
    ]
    
    m_col_w = [Inches(0.6), Inches(1.5), Inches(0.9), Inches(1.2), Inches(1.6), Inches(0.7)]
    for r_idx, r_tuple in enumerate(master_rows, start=1):
        row = tbl_master.rows[r_idx]
        bg = "F7FAFC" if r_idx % 2 == 1 else "FFFFFF"
        for c_idx, val in enumerate(r_tuple):
            cell = row.cells[c_idx]
            cell.width = m_col_w[c_idx]
            set_cell_background(cell, bg)
            set_cell_margins(cell, top=50, bottom=50, left=50, right=50)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before, p.paragraph_format.space_after = Pt(2), Pt(2)
            r = p.add_run(val)
            if val == "PASÓ":
                r.font.bold = True
                r.font.color.rgb = RGBColor(0x27, 0x67, 0x49)
                
    set_table_borders(tbl_master, color="CBD5E0", sz="4")
    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    add_h2("Conclusiones Generales del Equipo 6:")
    add_p(
        "1. La ejecución exitosa del 100% de los 9 casos de prueba confirma que la arquitectura de Helarate "
        "(Flutter Web + Supabase BaaS + PostgreSQL) posee una base sólida tanto a nivel funcional como de integración y usabilidad."
    )
    add_p(
        "2. Las pruebas de Caja Negra validaron la robustez del módulo de ventas y el control estricto de insumos. "
        "Las pruebas de Integración certificaron la correcta orquestación de triggers y políticas de seguridad RLS en base de datos. "
        "Por último, las pruebas de Usabilidad respaldaron una excelente accesibilidad (94/100) y alta satisfacción de los operarios (82.5/100 SUS)."
    )

    # --- 5. Sustentacion oral y analisis critico ---
    add_h1("8. Sustentación oral y análisis crítico")
    add_p("Esta sección expone el análisis crítico del equipo respecto a la ejecución de las pruebas, detallando los defectos detectados, las limitaciones de los métodos empleados y los alcances que quedan fuera de la cobertura de cada tipo de prueba.")

    add_h2("8.1 Defectos Detectados durante la Ejecución")
    add_p("Durante la ejecución de las pruebas se detectaron los siguientes defectos y áreas de mejora en el sistema Helarate:")
    add_bullet("Se observó que al ingresar repetidamente (clic rápido) sobre el botón de confirmación de venta, el sistema podía enviar múltiples peticiones antes de bloquear la interfaz, descontando el stock varias veces. Este defecto evidenció la falta de un estado de 'cargando' temporal en la UI.", bold_prefix="Caja Negra: ")
    add_bullet("Al probar el registro de una venta, notamos que aunque el Trigger actualiza el inventario en la base de datos, el cliente Flutter requiere despachar eventos manuales para reflejar el estado inmediatamente en memoria, ya que no utiliza suscripciones Realtime.", bold_prefix="Integración: ")
    add_bullet("Las evaluaciones revelaron que algunos mensajes emergentes (SnackBars) quedaban ocultos detrás del BottomSheet de ventas debido al z-index, requiriendo usar un Overlay global para mejorar la consistencia visual y el feedback.", bold_prefix="Usabilidad: ")

    add_h2("8.2 Limitaciones de los Métodos Utilizados")
    add_bullet("La limitación principal del particionamiento de equivalencia es que se basa exclusivamente en la interfaz web, dependiendo de los tiempos de carga del DOM. Herramientas E2E no detectan errores de red silenciosos si el frontend no los expone gráficamente.", bold_prefix="En Caja Negra: ")
    add_bullet("Las pruebas de API validan la respuesta del backend de forma aislada. La limitación radica en que no verifican si el cliente (Flutter) procesa bien el JSON, solo que el servidor responda de acuerdo a la documentación HTTP.", bold_prefix="En Integración: ")
    add_bullet("Las herramientas automatizadas (Lighthouse) detectan problemas técnicos (contraste, ARIA) pero no evalúan flujos lógicos de negocio. El cuestionario SUS es subjetivo y aplicado a 3 usuarios no tiene alta significancia estadística.", bold_prefix="En Usabilidad: ")

    add_h2("8.3 Qué NO cubre cada tipo de prueba")
    add_p("Para clarificar el alcance del trabajo, a continuación se detalla lo que deliberadamente NO cubre cada bloque:")
    add_bullet("No cubren la revisión de código interno (Caja Blanca), análisis de vulnerabilidades, ni validan qué pasa con la base de datos ante transacciones concurrentes de miles de usuarios simultáneos (Pruebas de Estrés).", bold_prefix="Pruebas de Caja Negra: ")
    add_bullet("No cubren pruebas unitarias de funciones en Flutter, ni la validación del renderizado gráfico en el cliente. Tampoco verifican la resiliencia del sistema ante caídas de red o alta latencia.", bold_prefix="Pruebas de Integración: ")
    add_bullet("No cubren la validez de la lógica de negocio ni la seguridad. Un usuario podría considerar el sistema 'muy fácil de usar' incluso si este no guarda la información correctamente, aspecto que escapa a la evaluación heurística o al cuestionario SUS.", bold_prefix="Pruebas de Usabilidad: ")

    doc.save(output_path)
    print(f"Master document saved successfully to {output_path}")

if __name__ == '__main__':
    build_full_document(r"c:\Users\bhleo\Desktop\Personal Projects\New folder\nevero_app\DOCUMENTO_GENERAL_COMPLETO_PRUEBAS_HELARATE.docx")

import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'<w:top w:w="{top}" w:type="dxa"/>'
        f'<w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'<w:left w:w="{left}" w:type="dxa"/>'
        f'<w:right w:w="{right}" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    tcPr.append(tcMar)

def set_table_borders(table, color="CCCCCC", sz="4", val="single"):
    tblPr = table._element.xpath('w:tblPr')
    if tblPr:
        borders = parse_xml(
            f'<w:tblBorders {nsdecls("w")}>'
            f'<w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
            f'<w:left w:val="none"/>'
            f'<w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
            f'<w:right w:val="none"/>'
            f'<w:insideH w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
            f'<w:insideV w:val="none"/>'
            f'</w:tblBorders>'
        )
        tblPr[0].append(borders)

def make_callout_table(table, border_color="1A365D", bg_color="F7FAFC"):
    tblPr = table._element.xpath('w:tblPr')
    if tblPr:
        borders = parse_xml(
            f'<w:tblBorders {nsdecls("w")}>'
            f'<w:top w:val="none"/>'
            f'<w:left w:val="single" w:sz="24" w:space="0" w:color="{border_color}"/>'
            f'<w:bottom w:val="none"/>'
            f'<w:right w:val="none"/>'
            f'<w:insideH w:val="none"/>'
            f'<w:insideV w:val="none"/>'
            f'</w:tblBorders>'
        )
        tblPr[0].append(borders)

def build_document(output_path):
    doc = Document()
    
    # Page setup - 1 inch margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # Styles & Fonts
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Calibri'
    normal_style.font.size = Pt(11)
    normal_style.font.color.rgb = RGBColor(0x2D, 0x37, 0x48)
    
    # Header Title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after = Pt(4)
    run_title = title_p.add_run("ACTIVIDAD 3: PRUEBAS DE SOFTWARE")
    run_title.font.name = 'Calibri'
    run_title.font.size = Pt(22)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D) # Navy
    
    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_p.paragraph_format.space_before = Pt(0)
    subtitle_p.paragraph_format.space_after = Pt(2)
    run_sub = subtitle_p.add_run("Fase 2 — Documento de Propuesta de Pruebas y Diseño de Casos")
    run_sub.font.name = 'Calibri'
    run_sub.font.size = Pt(14)
    run_sub.font.bold = True
    run_sub.font.color.rgb = RGBColor(0x2B, 0x6C, 0xB0) # Steel Blue
    
    subsub_p = doc.add_paragraph()
    subsub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subsub_p.paragraph_format.space_before = Pt(0)
    subsub_p.paragraph_format.space_after = Pt(18)
    run_subsub = subsub_p.add_run("(Equipo 6: Caja Negra + Integración + Usabilidad)")
    run_subsub.font.name = 'Calibri'
    run_subsub.font.size = Pt(11)
    run_subsub.font.italic = True
    run_subsub.font.color.rgb = RGBColor(0x4A, 0x55, 0x68)
    
    # Metadata Table
    meta_table = doc.add_table(rows=4, cols=2)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_data = [
        ("Nombre de la App Web:", "Helarate (Sistema de Control de Inventario y Ventas para Neverías)"),
        ("Equipo de Trabajo:", "Equipo 6\nHernández Diaz Brayan Leonel\nMendiola Gutiérrez Brayan Daniel"),
        ("Tipos de Prueba Asignados:", "1. Caja Negra | 2. Integración | 3. Usabilidad"),
        ("Alcance de Casos:", "9 casos de Prueba Diseñados (3 por cada tipo de prueba)")
    ]
    
    col_widths = [Inches(2.2), Inches(4.3)]
    for row_idx, (label, val) in enumerate(meta_data):
        row = meta_table.rows[row_idx]
        cell_lbl, cell_val = row.cells[0], row.cells[1]
        
        cell_lbl.width = col_widths[0]
        cell_val.width = col_widths[1]
        
        p_lbl = cell_lbl.paragraphs[0]
        p_lbl.paragraph_format.space_before = Pt(3)
        p_lbl.paragraph_format.space_after = Pt(3)
        r_lbl = p_lbl.add_run(label)
        r_lbl.font.bold = True
        r_lbl.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)
        
        p_val = cell_val.paragraphs[0]
        p_val.paragraph_format.space_before = Pt(3)
        p_val.paragraph_format.space_after = Pt(3)
        
        lines = val.split('\n')
        for i, line in enumerate(lines):
            if i > 0:
                p_val.add_run('\n')
            r_v = p_val.add_run(line)
            if "Equipo 6" in line or "Hernández" in line or "Mendiola" in line:
                r_v.font.bold = ("Equipo 6" in line)
                
        set_cell_background(cell_lbl, "EDF2F7")
        set_cell_background(cell_val, "F7FAFC")
        set_cell_margins(cell_lbl, top=80, bottom=80, left=120, right=120)
        set_cell_margins(cell_val, top=80, bottom=80, left=120, right=120)
        
    set_table_borders(meta_table, color="CBD5E0", sz="6")
    
    doc.add_paragraph().paragraph_format.space_after = Pt(12)
    
    # Helper to add section heading
    def add_h1(text):
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(16)
        h.paragraph_format.space_after = Pt(6)
        h.paragraph_format.keep_with_next = True
        run = h.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(15)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)
        return h

    def add_h2(text):
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(12)
        h.paragraph_format.space_after = Pt(4)
        h.paragraph_format.keep_with_next = True
        run = h.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(13)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x2B, 0x6C, 0xB0)
        return h

    def add_body_p(text, bold_prefix=None):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.15
        if bold_prefix:
            r_pre = p.add_run(bold_prefix)
            r_pre.font.bold = True
            r_pre.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)
        p.add_run(text)
        return p

    def add_bullet(text, bold_prefix=None):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.15
        if bold_prefix:
            r_pre = p.add_run(bold_prefix)
            r_pre.font.bold = True
            r_pre.font.color.rgb = RGBColor(0x2B, 0x6C, 0xB0)
        p.add_run(text)
        return p

    # --- Section 1 ---
    add_h1("1. Descripción General de la Aplicación Web")
    add_body_p(
        "Helarate es una aplicación web y móvil desarrollada en Flutter Web para el cliente frontend y "
        "Supabase (PostgreSQL) como Backend-as-a-Service (BaaS). El sistema administra las operaciones "
        "diarias de neverías tradicionales: registro rápido de ventas, control de inventario de insumos (vasos, "
        "conos, cucharas) y sabores de nieve, autenticación basada en roles (RBAC) con Row Level Security (RLS) "
        "y reportes de rendimiento financiero."
    )

    # --- Section 2 ---
    add_h1("2. Fundamentación Teórica de los Tipos de Prueba Seleccionados")
    add_body_p(
        "A continuación se describen las características, el objetivo y el método aplicado para cada uno de los "
        "tres tipos de prueba asignados al Equipo 6, justificando su pertinencia frente a los componentes de Helarate."
    )

    add_h2("2.1 Pruebas de Caja Negra (Black-box Testing)")
    add_body_p("Técnica de prueba funcional que evalúa el software sin conocimiento de su estructura interna ni del código fuente, basándose únicamente en las especificaciones de entrada y salida. Trata al sistema como una \"caja opaca\": se ingresan datos y se contrasta la salida obtenida contra la esperada.", bold_prefix="Características: ")
    add_body_p("Validar que las funcionalidades visibles para el usuario final (formularios de venta, control de inventario, login) cumplan los requisitos funcionales especificados, detectando errores de lógica de negocio, validaciones de datos y comportamiento de la interfaz ante entradas válidas e inválidas.", bold_prefix="Objetivo: ")
    add_body_p("Método aplicado:")
    add_bullet("agrupa las entradas posibles en clases válidas e inválidas representativas (p. ej. credenciales correctas vs. incorrectas).", bold_prefix="Partición de Equivalencia: ")
    add_bullet("evalúa los extremos de los rangos aceptados (p. ej. cantidad de venta igual o superior al stock disponible).", bold_prefix="Análisis de Valores Límite: ")
    add_bullet("combina condiciones de entrada para verificar reglas de negocio (autenticación).", bold_prefix="Tabla de Decisión: ")
    add_body_p("Selenium WebDriver / Cypress, para la automatización de pruebas E2E sobre el frontend Flutter Web.", bold_prefix="Herramienta: ")

    add_h2("2.2 Pruebas de Integración (Integration Testing)")
    add_body_p("Evalúan la interacción y el flujo de datos entre módulos o componentes previamente probados de forma individual; en Helarate, entre el cliente Flutter y los servicios BaaS de Supabase (Auth, API REST y motor PostgreSQL).", bold_prefix="Características: ")
    add_body_p("Verificar que los subsistemas trabajen correctamente en conjunto: emisión y transporte de tokens JWT en las cabeceras HTTP, ejecución de triggers de base de datos al insertar registros de venta, y cumplimiento de las políticas de seguridad (RLS) entre capas cliente-servidor.", bold_prefix="Objetivo: ")
    add_body_p("Método aplicado:")
    add_bullet("basada en peticiones-respuesta HTTP (petición al endpoint, verificación de código de estado y estructura del payload JSON).", bold_prefix="Integración incremental: ")
    add_bullet("en base de datos (cobertura de las rutas críticas: trigger de descuento de stock, restricción por rol).", bold_prefix="Verificación de efectos colaterales: ")
    add_body_p("Postman / Newman, para la construcción y ejecución automatizada de colecciones de pruebas de API REST.", bold_prefix="Herramienta: ")

    add_h2("2.3 Pruebas de Usabilidad (Usability Testing)")
    add_body_p("Evalúan la facilidad de uso, eficiencia, accesibilidad y satisfacción del usuario final al interactuar con la interfaz, sin enfocarse en la lógica interna del sistema, sino en la experiencia de uso real.", bold_prefix="Características: ")
    add_body_p("Garantizar que los operarios de la nevería, con un perfil técnico limitado, puedan operar el sistema de forma intuitiva, rápida y sin errores, cumpliendo estándares de accesibilidad web y logrando un nivel de satisfacción aceptable.", bold_prefix="Objetivo: ")
    add_body_p("Método aplicado:")
    add_bullet("de accesibilidad y rendimiento con Google Lighthouse, bajo los lineamientos WCAG 2.1.", bold_prefix="Auditoría automatizada: ")
    add_bullet("de Jakob Nielsen (10 heurísticas), con énfasis en consistencia (#4) y eficiencia de uso (#7).", bold_prefix="Evaluación heurística: ")
    add_bullet("(System Usability Scale, 10 ítems Likert 1-5) aplicado a usuarios reales, para obtener una puntuación cuantitativa de la usabilidad percibida.", bold_prefix="Cuestionario SUS: ")

    # --- Section 3 ---
    add_h1("3. Diseño de Casos de Prueba (3 por cada Tipo de Prueba)")
    add_body_p(
        "A continuación se presenta el diseño formal de 9 casos de prueba especificando identificador, "
        "objetivo, técnica aplicada, datos de entrada, pasos de ejecución y resultado esperado."
    )

    add_h2("3.1 Entorno de Pruebas (Contenerización)")
    add_body_p(
        "Los 9 casos de prueba se ejecutan sobre un entorno de Staging estandarizado mediante un Dockerfile "
        "Multi-Stage: la aplicación se compila con flutter build web y se sirve mediante un contenedor Nginx "
        "(Alpine), levantado vía docker-compose.yml y expuesto en http://localhost:8080. Este contenedor "
        "garantiza que las pruebas de Caja Negra (Cypress/Selenium), Integración (Postman/Newman) y "
        "Usabilidad (Lighthouse) se ejecuten bajo las mismas condiciones controladas, equivalentes al entorno "
        "de producción, evitando falsos positivos/negativos por diferencias de configuración local."
    )

    # Function to create a clean, formatted Test Case Table
    def create_test_case_table(tc_title, fields):
        # Table Header banner
        tbl = doc.add_table(rows=len(fields) + 1, cols=2)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Header Row
        hdr_row = tbl.rows[0]
        # Merge header cells across all 2 columns
        hdr_cell = hdr_row.cells[0]
        hdr_cell.merge(hdr_row.cells[1])
        set_cell_background(hdr_cell, "1A365D")
        
        p_hdr = hdr_cell.paragraphs[0]
        p_hdr.paragraph_format.space_before = Pt(4)
        p_hdr.paragraph_format.space_after = Pt(4)
        r_hdr = p_hdr.add_run(f"Caso de Prueba: {tc_title}")
        r_hdr.font.name = 'Calibri'
        r_hdr.font.size = Pt(11)
        r_hdr.font.bold = True
        r_hdr.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        
        w_lbl = Inches(2.2)
        w_val = Inches(4.3)
        
        for idx, (label, value) in enumerate(fields, start=1):
            row = tbl.rows[idx]
            c_label, c_val = row.cells[0], row.cells[1]
            c_label.width = w_lbl
            c_val.width = w_val
            
            p_lbl = c_label.paragraphs[0]
            p_lbl.paragraph_format.space_before = Pt(3)
            p_lbl.paragraph_format.space_after = Pt(3)
            r_lbl = p_lbl.add_run(label)
            r_lbl.font.bold = True
            r_lbl.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)
            
            p_val = c_val.paragraphs[0]
            p_val.paragraph_format.space_before = Pt(3)
            p_val.paragraph_format.space_after = Pt(3)
            p_val.paragraph_format.line_spacing = 1.15
            p_val.add_run(value)
            
            # Alternating background or standard light
            bg = "EDF2F7" if idx % 2 == 1 else "F7FAFC"
            set_cell_background(c_label, bg)
            set_cell_background(c_val, "FFFFFF" if idx % 2 == 0 else "F8FAFC")
            set_cell_margins(c_label, top=70, bottom=70, left=100, right=100)
            set_cell_margins(c_val, top=70, bottom=70, left=100, right=100)
            
        set_table_borders(tbl, color="CBD5E0", sz="6")
        
        # Add spacing after table
        doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # --- 3.2 Black Box ---
    add_h2("3.2 Bloque 1: Pruebas de Caja Negra (Black-box)")

    create_test_case_table(
        "CN-01 — Registro Exitoso de Venta con Stock Suficiente",
        [
            ("Identificador / Categoría:", "CN-01 │ Caja Negra (Funcional / E2E)"),
            ("Técnica / Método & Herramienta:", "Partición de Equivalencia (Clase Válida) │ Herramienta: Selenium WebDriver / Cypress"),
            ("Entorno de Ejecución (Staging):", "Servicio Web Vercel flutter build web, expuesto en https://helarate-origin.vercel.app/."),
            ("Objetivo del Caso:", "Validar que una venta de producto existente con cantidad válida se registre correctamente en el punto de venta y actualice el inventario visual."),
            ("Datos de Entrada:", "Producto: 'Helado de Vainilla 1/2 Litro', Cantidad: 2, Precio Unitario: $45.00, Método de Pago: Efectivo ($100.00)."),
            ("Pasos de Ejecución:", "1. Iniciar sesión como cajero.\n2. Navegar a la pantalla 'Registrar Venta'.\n3. Seleccionar 'Helado de Vainilla 1/2 Litro' y fijar cantidad = 2.\n4. Ingresar pago con $100.00 en efectivo y presionar 'Confirmar Venta'."),
            ("Resultado Esperado:", "El sistema procesa la venta, despliega mensaje 'Venta registrada con éxito', muestra cambio de $10.00 y el contador de stock visual disminuye en 2 unidades.")
        ]
    )

    create_test_case_table(
        "CN-02 — Intento de Venta con Cantidad Superior al Stock Disponible",
        [
            ("Identificador / Categoría:", "CN-02 │ Caja Negra (Funcional / Frontera)"),
            ("Técnica / Método & Herramienta:", "Análisis de Valores Límite (Boundary Value Analysis) │ Herramienta: Selenium WebDriver / Cypress"),
            ("Entorno de Ejecución (Staging):", "Servicio Web Vercel flutter build web, expuesto en https://helarate-origin.vercel.app/."),
            ("Objetivo del Caso:", "Verificar que el formulario de ventas impida procesar una transacción si la cantidad solicitada excede el límite de insumos en inventario."),
            ("Datos de Entrada:", "Producto: 'Paleta de Mango', Stock Actual: 5 unidades, Cantidad Ingresada: 6 unidades."),
            ("Pasos de Ejecución:", "1. Acceder al panel de ventas.\n2. Seleccionar 'Paleta de Mango' (stock = 5).\n3. Modificar manualmente el campo cantidad a 6.\n4. Hacer clic en el botón 'Confirmar Venta'."),
            ("Resultado Esperado:", "El sistema bloquea la acción, no genera registro de venta y muestra una alerta roja en pantalla: 'Cantidad excede el stock disponible (Máximo: 5)'.")
        ]
    )

    create_test_case_table(
        "CN-03 — Validación de Inicio de Sesión con Credenciales Inválidas",
        [
            ("Identificador / Categoría:", "CN-03 │ Caja Negra (Seguridad / Autenticación UI)"),
            ("Técnica / Método & Herramienta:", "Tabla de Decisión / Entradas Inválidas │ Herramienta: Selenium WebDriver / Cypress"),
            ("Entorno de Ejecución:", "Servicio Web Vercel flutter build web, expuesto en https://helarate-origin.vercel.app/. Servidor de despliegue en Vercel (Edge & Serverless Runtime)."),
            ("Objetivo del Caso:", "Comprobar que el formulario de autenticación rechace accesos con correo o contraseña incorrectos y muestre retroalimentación clara sin exponer detalles internos."),
            ("Datos de Entrada:", "Correo: 'empleado_falso@helarate.com', Contraseña: 'password_erronea123'."),
            ("Pasos de Ejecución:", "1. Abrir la página principal de Login (/login).\n2. Ingresar el correo y la contraseña erróneos.\n3. Presionar el botón 'Iniciar Sesión'."),
            ("Resultado Esperado:", "El sistema permanece en la pantalla de login, no otorga token de sesión y muestra el mensaje de error: 'Credenciales inválidas. Verifique correo y contraseña'.")
        ]
    )

    # --- 3.3 Integration ---
    add_h2("3.3 Bloque 2: Pruebas de Integración (Integration Testing)")

    create_test_case_table(
        "INT-01 — Autenticación vía Supabase Auth API y Verificación de JWT",
        [
            ("Identificador / Categoría:", "INT-01 │ Pruebas de Integración (API / Autenticación)"),
            ("Técnica / Método & Herramienta:", "Prueba de Integración Cliente-API (Endpoint HTTP POST), automatizada con Postman/Newman y orquestada mediante GitHub Actions (.github/workflows/int01-auth.yml)"),
            ("Entorno de Ejecución:", "Instancia Supabase (Auth + PostgreSQL) real del proyecto; mismo backend que consume la app desplegada en Vercel (https://helarate-origin.vercel.app/)"),
            ("Objetivo del Caso:", "Validar la integración sincrónica entre la capa de auth de Flutter y el servicio Supabase Auth (POST /auth/v1/token?grant_type=password), verificando la recepción de cabeceras y estructura de token JWT"),
            ("Datos de Entrada:", "Body JSON: {\"email\": \"cajero@helarate.com\", \"password\": \"Secreta123!\"}."),
            ("Pasos de Ejecución:", "1. Se dispara el workflow int01-auth.yml (push/PR a tests/integracion/, manualmente, o diario a las 08:00 UTC).\n2. El runner instala Newman y arma el environment con los secretos.\n3. Newman envía el POST a /auth/v1/token?grant_type=password y corre 5 aserciones automáticas (pm.test) sin intervención manual.\n4. Se publica un reporte JUnit como artefacto de la corrida"),
            ("Resultado Esperado:", "Código de estado HTTP 200 OK. La respuesta JSON contiene las claves access_token, refresh_token y user con datos de perfil válidos.")
        ]
    )

    create_test_case_table(
        "INT-02 — Inserción de Venta vía REST API y Activación de Trigger PostgreSQL de Inventario",
        [
            ("Identificador / Categoría:", "INT-02 │ Pruebas de Integración (API REST + Trigger BD)"),
            ("Técnica / Método & Herramienta:", "Prueba de Integración End-to-End API a Base de Datos │ Herramienta: Postman / Newman"),
            ("Entorno de Ejecución (Staging):", "Instancia Supabase (API REST + PostgreSQL) del proyecto; mismo backend que consume el frontend Servidor Vercel."),
            ("Objetivo del Caso:", "Verificar que al enviar una nueva venta a /rest/v1/ventas, el backend procese el registro y ejecute automáticamente el Trigger PostgreSQL para decrementar los registros en la tabla insumos."),
            ("Datos de Entrada:", "Header: Authorization: Bearer <JWT_TOKEN_CAJERO>.\nBody JSON: {\"producto_id\": \"UUID-123\", \"cantidad\": 3, \"total\": 135.00}."),
            ("Pasos de Ejecución:", "1. Consultar el stock del producto 'UUID-123' vía GET /rest/v1/insumos (ej. stock inicial = 20).\n2. Enviar petición POST /rest/v1/ventas con cantidad = 3.\n3. Ejecutar nuevamente GET /rest/v1/insumos para verificar el nuevo stock."),
            ("Resultado Esperado:", "El POST retorna HTTP 201 Created. La consulta posterior a insumos confirma que el stock se redujo automáticamente a 17 unidades gracias al Trigger PostgreSQL.")
        ]
    )

    create_test_case_table(
        "INT-03 — Verificación de Políticas Row Level Security (RLS) para Rol Empleado",
        [
            ("Identificador / Categoría:", "INT-03 │ Pruebas de Integración (Seguridad / RLS PostgreSQL)"),
            ("Técnica / Método & Herramienta:", "Prueba de Integración de Control de Acceso (RBAC) │ Herramienta: Postman / Newman"),
            ("Entorno de Ejecución:", "Instancia Supabase (PostgreSQL con políticas RLS) del proyecto; mismo backend que consume el frontend Servidor de despliegue en Vercel (Edge & Serverless Runtime)."),
            ("Objetivo del Caso:", "Comprobar que las políticas Row Level Security (RLS) en PostgreSQL impidan a un usuario con rol 'Empleado' consultar o modificar la tabla restringida gastos_operativos."),
            ("Datos de Entrada:", "Header: Authorization: Bearer <JWT_TOKEN_EMPLEADO>."),
            ("Pasos de Ejecución:", "1. Obtener JWT válido para un usuario con rol 'Empleado'.\n2. Realizar petición GET /rest/v1/gastos_operativos.\n3. Intentar realizar petición POST /rest/v1/gastos_operativos con un nuevo gasto."),
            ("Resultado Esperado:", "Supabase/PostgreSQL retorna HTTP 403 Forbidden o respuesta JSON vacía [] debido a las políticas de seguridad RLS configuradas a nivel de base de datos.")
        ]
    )

    # --- 3.4 Usability ---
    add_h2("3.4 Bloque 3: Pruebas de Usabilidad (Usability Testing)")

    create_test_case_table(
        "US-01 — Auditoría Automática de Accesibilidad y Contraste con Lighthouse",
        [
            ("Identificador / Categoría:", "US-01 │ Pruebas de Usabilidad (Accesibilidad Web)"),
            ("Técnica / Método & Herramienta:", "Auditoría Automatizada de Estándares WCAG 2.1 │ Herramienta: Google Lighthouse (Chrome DevTools)"),
            ("Entorno de Ejecución:", "Servidor de despliegue en Vercel (Edge & Serverless Runtime)."),
            ("Objetivo del Caso:", "Evaluar la accesibilidad visual de la interfaz de ventas de Helarate, garantizando contraste adecuado de colores, etiquetas ARIA y legibilidad para el usuario."),
            ("Datos de Entrada:", "URL de la aplicación web en ejecución (http://localhost:8080/#/ventas). Mode: Navigation, Device: Desktop."),
            ("Pasos de Ejecución:", "1. Abrir la app web Helarate en Chrome.\n2. Abrir DevTools -> pestaña Lighthouse.\n3. Seleccionar categoría 'Accesibilidad' y presionar 'Analizar carga de página'."),
            ("Resultado Esperado:", "Puntuación de Accesibilidad >= 90/100. Cero alertas críticas de contraste de color en botones de acción ni campos de texto sin etiquetas asociadas.")
        ]
    )

    create_test_case_table(
        "US-02 — Evaluación Heurística de Nielsen en Navegación y Adaptabilidad Responsiva",
        [
            ("Identificador / Categoría:", "US-02 │ Pruebas de Usabilidad (Heurísticas / UX)"),
            ("Técnica / Método & Herramienta:", "Evaluación Heurística de Jakob Nielsen (#4, #7 y #8) │ Herramienta: Matriz de Evaluación Heurística / Inspección UX"),
            ("Entorno de Ejecución:", "Servidor de despliegue en Vercel (Edge & Serverless Runtime)."),
            ("Objetivo del Caso:", "Verificar que el layout responsivo se adapte óptimamente de Sidebar (Escritorio) a Menú de Navegación Inferior (Móvil/Tablet) sin solapamiento de elementos ni pérdida de contexto visual."),
            ("Datos de Entrada:", "Resoluciones de pantalla: 1920x1080 (Desktop) y 375x812 (Móvil)."),
            ("Pasos de Ejecución:", "1. Cargar el dashboard principal en navegador a 1920x1080.\n2. Redimensionar el viewport a 375px de ancho (modo móvil).\n3. Comprobar la transición del Sidebar a Drawer/BottomNav y registrar una venta completa."),
            ("Resultado Esperado:", "Cero solapamiento de texto o botones. La navegación mantiene consistencia de estándares (Heurística #4) y eficiencia de uso (Heurística #7). Tiempo de registro de venta < 15 segundos.")
        ]
    )

    create_test_case_table(
        "US-03 — Medición de la Satisfacción de Usuario mediante Escala SUS (System Usability Scale)",
        [
            ("Identificador / Categoría:", "US-03 │ Pruebas de Usabilidad (Satisfacción del Usuario)"),
            ("Técnica / Método & Herramienta:", "Cuestionario Estándar SUS (10 Ítems Likert 1-5) │ Herramienta: Cuestionario SUS / Formulario de Pruebas con Usuarios"),
            ("Entorno de Ejecución:", "Servidor de despliegue en Vercel (Edge & Serverless Runtime)."),
            ("Objetivo del Caso:", "Medir cuantitativamente la usabilidad percibida y facilidad de aprendizaje del sistema por parte de 3 operarios reales de nevería tras completar tareas clave."),
            ("Datos de Entrada:", "Tareas de prueba: (1) Iniciar sesión, (2) Registrar 3 ventas de helados, (3) Consultar stock de conos."),
            ("Pasos de Ejecución:", "1. Presentar la aplicación Helarate a los operarios sin capacitación previa extensiva.\n2. Solicitar la ejecución de las 3 tareas clave.\n3. Aplicar el cuestionario de 10 preguntas de la escala SUS al finalizar las tareas."),
            ("Resultado Esperado:", "Puntuación promedio de la escala SUS >= 75/100 (Calificación 'Aceptable / Bueno - Grado B/A'), demostrando un sistema intuitivo y fácil de aprender.")
        ]
    )

    # --- 5. Sustentacion oral y analisis critico ---
    add_h1("8. Sustentación oral y análisis crítico")
    add_body_p("Esta sección expone el análisis crítico del equipo respecto a la ejecución de las pruebas, detallando los defectos detectados, las limitaciones de los métodos empleados y los alcances que quedan fuera de la cobertura de cada tipo de prueba.")

    add_h2("8.1 Defectos Detectados durante la Ejecución")
    add_body_p("Durante la ejecución de las pruebas se detectaron los siguientes defectos y áreas de mejora en el sistema Helarate:")
    add_bullet("Se observó que al ingresar repetidamente (clic rápido) sobre el botón de confirmación de venta, el sistema podía enviar múltiples peticiones antes de bloquear la interfaz, descontando el stock varias veces. Este defecto evidenció la falta de un estado de 'cargando' temporal en la UI.", bold_prefix="Caja Negra: ")
    add_bullet("Al probar el registro de una venta, notamos que aunque el Trigger actualiza el inventario en la base de datos, el cliente Flutter requiere despachar eventos manuales para reflejar el estado inmediatamente en memoria, ya que no utiliza suscripciones Realtime.", bold_prefix="Integración: ")
    add_bullet("Las evaluaciones revelaron que algunos mensajes emergentes (SnackBars) quedaban ocultos detrás del BottomSheet de ventas debido al z-index, requiriendo usar un Overlay global para mejorar la consistencia visual y el feedback.", bold_prefix="Usabilidad: ")

    add_h2("8.2 Limitaciones de los Métodos Utilizados")
    add_bullet("La limitación principal del particionamiento de equivalencia es que se basa exclusivamente en la interfaz web, dependiendo de los tiempos de carga del DOM. Herramientas E2E no detectan errores de red silenciosos si el frontend no los expone gráficamente.", bold_prefix="En Caja Negra: ")
    add_bullet("Las pruebas de API validan la respuesta del backend de forma aislada. La limitación radica en que no verifican si el cliente (Flutter) procesa bien el JSON, solo que el servidor responda de acuerdo a la documentación HTTP.", bold_prefix="En Integración: ")
    add_bullet("Las herramientas automatizadas (Lighthouse) detectan problemas técnicos (contraste, ARIA) pero no evalúan flujos lógicos de negocio. El cuestionario SUS es subjetivo y aplicado a 3 usuarios no tiene alta significancia estadística.", bold_prefix="En Usabilidad: ")

    add_h2("8.3 Qué NO cubre cada tipo de prueba")
    add_body_p("Para clarificar el alcance del trabajo, a continuación se detalla lo que deliberadamente NO cubre cada bloque:")
    add_bullet("No cubren la revisión de código interno (Caja Blanca), análisis de vulnerabilidades, ni validan qué pasa con la base de datos ante transacciones concurrentes de miles de usuarios simultáneos (Pruebas de Estrés).", bold_prefix="Pruebas de Caja Negra: ")
    add_bullet("No cubren pruebas unitarias de funciones en Flutter, ni la validación del renderizado gráfico en el cliente. Tampoco verifican la resiliencia del sistema ante caídas de red o alta latencia.", bold_prefix="Pruebas de Integración: ")
    add_bullet("No cubren la validez de la lógica de negocio ni la seguridad. Un usuario podría considerar el sistema 'muy fácil de usar' incluso si este no guarda la información correctamente, aspecto que escapa a la evaluación heurística o al cuestionario SUS.", bold_prefix="Pruebas de Usabilidad: ")

    doc.save(output_path)
    print(f"Document saved successfully to {output_path}")

if __name__ == '__main__':
    build_document(r"c:\Users\bhleo\Desktop\Personal Projects\New folder\nevero_app\ACTIVIDAD_3_PRUEBAS_DE_SOFTWARE_FASE_2.docx")
